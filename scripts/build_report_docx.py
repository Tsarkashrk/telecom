from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = ROOT / "docs"
OUTPUT_PATH = OUTPUT_DIR / "Практическая_работа_6_Отчет_по_заданию_1.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
TEXT = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
LIGHT_FILL = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, inches):
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(inches * 1440)))


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "C9D2DC")
        borders.append(border)
    tbl_pr.append(borders)


def set_run_font(run, *, size=11, bold=False, color=TEXT, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def style_paragraph(paragraph, *, before=0, after=6, line=1.1, align=WD_ALIGN_PARAGRAPH.LEFT):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    paragraph.alignment = align


def add_body(document, text):
    p = document.add_paragraph()
    style_paragraph(p)
    set_run_font(p.add_run(text))
    return p


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        style_paragraph(p, after=4, line=1.15)
        set_run_font(p.add_run(item))


def add_numbered(document, items):
    for item in items:
        p = document.add_paragraph(style="List Number")
        style_paragraph(p, after=4, line=1.15)
        set_run_font(p.add_run(item))


def add_heading(document, text, level):
    p = document.add_paragraph()
    style_paragraph(p, before=16 if level == 1 else 12, after=6)
    run = p.add_run(text)
    set_run_font(
        run,
        size=16 if level == 1 else 13,
        bold=True,
        color=BLUE if level == 1 else DARK_BLUE,
    )
    return p


def create_doc():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = document.add_paragraph()
    style_paragraph(title, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(title.add_run("ПРАКТИЧЕСКАЯ РАБОТА № 6"), size=18, bold=True, color=TEXT)

    subtitle = document.add_paragraph()
    style_paragraph(subtitle, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(
        subtitle.add_run(
            "Разработка и оценка защищенности MVP-продукта в соответствии с принципами Secure SDLC и OWASP Top 10"
        ),
        size=13,
        bold=True,
        color=DARK_BLUE,
    )

    meta = [
        ("Тема", "Разработка защищенного MVP телеком-платформы с пользовательским интерфейсом, биллингом и аудитом безопасности."),
        ("Цель работы", "Сформировать практические навыки проектирования, реализации, тестирования и аудита защищенного программного продукта."),
        ("Вариант", "Телекоммуникационная платформа регистрации клиентов, подключения тарифов, выставления и оплаты счетов."),
    ]

    for label, value in meta:
        p = document.add_paragraph()
        style_paragraph(p, after=4)
        set_run_font(p.add_run(f"{label}: "), bold=True)
        set_run_font(p.add_run(value))

    add_body(
        document,
        "Проект реализован на основе существующего backend на FastAPI и дополнен новым frontend на React + Vite. Решение охватывает основной бизнес-процесс подключения абонента к тарифу, оплаты счета и активации услуги с учетом ролевой модели, валидации и журналирования критичных действий.",
    )

    add_heading(document, "1. Краткое описание разработанного MVP", 1)
    add_body(
        document,
        "Разработанный MVP представляет собой клиент-серверную информационную систему для телеком-оператора. Пользователь может зарегистрироваться, войти в личный кабинет, выбрать тарифный план, активировать подписку в режиме предоплаты, просмотреть счета и оплатить их. После оплаты подписка переводится в активное состояние. Для операторов и администраторов предусмотрен контролируемый доступ к данным абонентов.",
    )
    add_bullets(
        document,
        [
            "Backend: FastAPI, SQLAlchemy, JWT-аутентификация, аудит безопасности.",
            "Frontend: React + Vite, адаптивный интерфейс, панель абонента и оператора.",
            "База данных: сущности User, TariffPlan, Subscription, Invoice, AuditLog.",
            "Основные роли: customer, operator, admin.",
        ],
    )

    add_heading(document, "2. Архитектура системы", 1)
    add_numbered(
        document,
        [
            "Frontend SPA на React + Vite взаимодействует с backend по HTTP API.",
            "Backend на FastAPI обрабатывает аутентификацию, подписки, биллинг и внутренние сервисные операции.",
            "SQLAlchemy реализует слой доступа к данным и работу с сущностями базы данных.",
            "Подсистема безопасности отвечает за проверку входных данных, JWT, RBAC, аудит и безопасную обработку ошибок.",
            "Журнал аудита фиксирует входы в систему, ошибки доступа, создание и оплату счетов, а также иные критичные события.",
        ],
    )

    add_heading(document, "3. Роли пользователей и разграничение доступа", 1)
    add_bullets(
        document,
        [
            "customer: регистрация, вход, просмотр собственных тарифов, подписок и счетов, оплата только собственных счетов.",
            "operator: возможности просмотра подписок и счетов других пользователей через специально защищенные endpoint’ы.",
            "admin: все возможности оператора, полный привилегированный доступ в рамках реализованных API.",
        ],
    )
    add_body(
        document,
        "Разграничение доступа реализовано в backend через проверки ролей и принадлежности ресурса владельцу. Доступ к чужим подпискам и счетам для обычного клиента запрещен и фиксируется как событие безопасности.",
    )

    add_heading(document, "4. Основной бизнес-сценарий", 1)
    add_numbered(
        document,
        [
            "Пользователь проходит регистрацию.",
            "Пользователь выполняет вход и получает access token и refresh token.",
            "Пользователь открывает каталог тарифов и выбирает нужный план.",
            "Система создает подписку в статусе pending_payment и автоматически формирует счет.",
            "Пользователь переходит в раздел биллинга и оплачивает счет.",
            "После оплаты backend активирует подписку, обновляет статус счета и сохраняет события в журнал аудита.",
        ],
    )

    add_heading(document, "5. Перечень реализованных механизмов безопасности", 1)
    add_bullets(
        document,
        [
            "Хеширование паролей с использованием bcrypt/passlib.",
            "JWT access token и refresh token с проверкой подписи, срока жизни и типа токена.",
            "Ротация refresh token и защита от повторного использования старого токена.",
            "Ограничение brute-force попыток входа и временная блокировка после серии ошибок.",
            "Проверка и нормализация входных данных через Pydantic и дополнительные функции безопасности.",
            "Ограничение размера HTTP request body middleware-слоем.",
            "RBAC и owner-based authorization для счетов и подписок.",
            "Санитизация CSV-экспорта и проверка безопасного имени файла.",
            "Журналирование критичных действий без утечки секретов, токенов и SQL-ошибок наружу.",
            "Хранение секретов и параметров подключения во внешнем .env-файле.",
        ],
    )

    add_heading(document, "6. Результаты анализа по OWASP Top 10", 1)
    add_body(
        document,
        "При анализе защищенности были рассмотрены основные категории риска OWASP Top 10, применимые к данному проекту. Проверка показала, что для большинства критичных рисков в системе уже реализованы компенсирующие меры либо внесены исправления в код и конфигурацию.",
    )
    add_bullets(
        document,
        [
            "Broken Access Control: введены проверки принадлежности ресурса и привилегированных ролей.",
            "Authentication Failures: реализованы lockout после неудачных попыток и refresh token rotation.",
            "Injection: используется SQLAlchemy ORM, входные данные валидируются и нормализуются.",
            "Security Misconfiguration: секреты вынесены из кода, предусмотрен .env.example.",
            "Cryptographic Failures: пароли не хранятся открыто, токены подписываются и проверяются.",
            "Software Supply Chain Failures: выполнена проверка frontend-зависимостей через npm audit.",
            "Software or Data Integrity Failures: введена проверка версии refresh token и подписи JWT.",
            "Security Logging and Alerting Failures: аудит охватывает вход, отказ, создание счета, оплату и несанкционированный доступ.",
            "Mishandling of Exceptional Conditions: backend возвращает безопасные сообщения об ошибках без технических подробностей.",
        ],
    )

    add_heading(document, "7. Таблица выявленных уязвимостей и мер по их устранению", 1)
    table = document.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    widths = [1.2, 1.2, 2.0, 1.55, 0.8, 1.75]
    headers = [
        "Категория риска",
        "Место обнаружения",
        "Описание уязвимости",
        "Последствия",
        "Критичность",
        "Способ исправления",
    ]
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[idx])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_run_font(p.add_run(headers[idx]), size=10, bold=True)

    rows = [
        [
            "Broken Access Control",
            "subscriptions, invoices",
            "Риск доступа клиента к чужим счетам или подпискам по идентификатору.",
            "Утечка данных и несанкционированные действия.",
            "Высокая",
            "Добавлены проверки ensure_subscription_access и ensure_invoice_access.",
        ],
        [
            "Authentication Failures",
            "auth/login, auth/refresh",
            "Риск перебора пароля и повторного использования старого refresh token.",
            "Захват учетной записи.",
            "Высокая",
            "Введены lockout, ротация refresh token и проверка token version.",
        ],
        [
            "Injection",
            "ORM-запросы, CSV export",
            "Потенциальные SQL- и CSV-injection при небезопасной обработке данных.",
            "Порча данных, выполнение формул в табличных редакторах.",
            "Средняя",
            "Использованы ORM, валидация input и csv_sanitize_cell.",
        ],
        [
            "Security Misconfiguration",
            "config, .env",
            "Риск хранения секретов и параметров подключения в исходном коде.",
            "Компрометация ключей и инфраструктуры.",
            "Высокая",
            "Секреты вынесены во внешнюю конфигурацию и .env.example.",
        ],
        [
            "Cryptographic Failures",
            "security.py",
            "Опасность хранения паролей без хеширования или слабой проверки токенов.",
            "Компрометация учетных данных.",
            "Высокая",
            "Использованы bcrypt/passlib и проверка подписи JWT.",
        ],
        [
            "Software Supply Chain Failures",
            "frontend/package.json",
            "Уязвимые dev-зависимости Vite/esbuild по результатам первичного npm audit.",
            "Локальные риски при разработке и dev-server эксплуатации.",
            "Средняя",
            "Зависимости обновлены до безопасных версий, повторный audit показал 0 vulnerabilities.",
        ],
        [
            "Logging and Monitoring Failures",
            "logging_config.py",
            "Недостаточное журналирование критичных событий безопасности.",
            "Снижение наблюдаемости инцидентов.",
            "Средняя",
            "Расширен аудит входов, отказов, платежей и unauthorized access attempt.",
        ],
        [
            "Mishandling of Exceptional Conditions",
            "main.py",
            "Риск возврата пользователю технических деталей внутренних ошибок.",
            "Упрощение разведки приложения.",
            "Средняя",
            "Реализованы безопасные глобальные exception handler’ы.",
        ],
    ]

    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row.cells):
            set_cell_width(value, widths[idx])
            value.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = value.paragraphs[0]
            style_paragraph(p, after=0, line=1.05)
            set_run_font(p.add_run(row_data[idx]), size=9)

    document.add_paragraph()

    add_heading(document, "8. Результаты повторного тестирования", 1)
    add_bullets(
        document,
        [
            "Backend-тесты: выполнен запуск ./venv/bin/pytest, получено 28 passed.",
            "Проверены регистрация, аутентификация, валидация полей, ограничения доступа, export, refresh token rotation и oversized request body.",
            "Frontend: выполнена production-сборка npm run build без ошибок.",
            "Проверка цепочки поставок: после обновления зависимостей выполнен npm audit, получено 0 vulnerabilities.",
        ],
    )
    add_body(
        document,
        "Повторное тестирование подтвердило корректность основного бизнес-сценария и устранение выявленного риска в dev-цепочке frontend-зависимостей.",
    )

    add_heading(document, "9. Выводы по работе", 1)
    add_body(
        document,
        "В результате выполнения практической работы был сформирован завершенный MVP-продукт с backend и современным frontend-интерфейсом. Проект реализует основной бизнес-процесс телеком-платформы, включает аутентификацию, разграничение доступа, журналирование и базовые механизмы защищенной разработки. Анализ по OWASP Top 10 позволил выявить и устранить ряд рисков, а повторное тестирование подтвердило корректность и повышенную защищенность системы.",
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(footer, after=0)
    set_run_font(footer.add_run("Telecom Secure MVP"), size=9, color=MUTED)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = create_doc()
    print(path)
